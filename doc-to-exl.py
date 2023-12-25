from docx import Document
from openpyxl import Workbook
from tkinter import Tk, filedialog
import re
import time



def parse_docx_to_excel(docx_filename, excel_filename):
    time_start = time.time()
    doc = Document(docx_filename)

    wb = Workbook()
    ws = wb.active

    first_line = doc.paragraphs[0].text if doc.paragraphs else ""
    ws['A1'] = first_line

    if len(doc.paragraphs) > 1:
        for paragraph in doc.paragraphs:
            text = paragraph.text
            pattern_one = "зат. геом."
            pattern_two = "/*"
            pattern_three = "(Смирнов"
            parts = text.split(pattern_one)
            termin = parts[0] if len(parts) > 0 else ""
            field = pattern_one
            definition = ''
            example = ''
            book = ''
            page = 'ERROR'
            if len(parts) > 1:
                parts_two = parts[1].split(pattern_two)
                definition = parts_two[0]
                if len(parts_two) > 1:
                    parts_three = parts_two[1].split(pattern_three)
                    example = parts_three[0]
                else:
                    parts_three = parts[1].split(pattern_three)
                    definition = parts_three[0]
                match_double = re.search(r"(\d+)-(\d+) б\.|(\d+)-(\d+)б\.", text)
                match = re.search(r"(\d+) б\.|(\d+)б\.", text)
                if match_double:
                    number = f"{match_double.group(1)}-{match_double.group(2)}"
                    page = number
                elif match:
                    number = match.group(1)
                    page = number
                
                book = ''
            end = ''

            row_data = ['',termin, field, definition, example, page, book, end]
            ws.append(row_data)

    wb.save(excel_filename)
    time_finish = time.time()
    print(f"Файл {excel_filename} был обработан за {round(time_finish - time_start, 2)}, секунд!")



def open_file_dialog():
    root = Tk()
    root.withdraw()
    output_excel_file = ''
    
    file_path = filedialog.askopenfilename(title="WORD DOC", filetypes=(("Word files", "*.docx"), ("All files", "*.*")))
    
    if file_path:
        output_excel_file = f'{file_path}.xlsx'
        parse_docx_to_excel(file_path, output_excel_file)


open_file_dialog()
