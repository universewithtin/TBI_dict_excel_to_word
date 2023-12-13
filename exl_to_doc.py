from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
import time
import re
import sys

file_extension = ".xlsx"

def clean_end_of_line(text, add_dot=True):
    text = re.sub(r'[. ]*$', '', text)
    if add_dot:
        if text.endswith('?'):
            text += ' '
        elif text.endswith(';'):
            text = text[:-1] + '. '
        else:
            text += '. '
    else:
        text += ' '
    return text


def run_set_spacing(run, value: int):
    """Set the font spacing for `run` to `value` in twips.

    A twip is a "twentieth of an imperial point", so 1/1440 in.
    """

    def get_or_add_spacing(rPr):
        # --- check if `w:spacing` child already exists ---
        spacings = rPr.xpath("./w:spacing")
        # --- return that if so ---
        if spacings:
            return spacings[0]
        # --- otherwise create one ---
        spacing = OxmlElement("w:spacing")
        rPr.insert_element_before(
            spacing,
            *(
                "w:w",
                "w:kern",
                "w:position",
                "w:sz",
                "w:szCs",
                "w:highlight",
                "w:u",
                "w:effect",
                "w:bdr",
                "w:shd",
                "w:fitText",
                "w:vertAlign",
                "w:rtl",
                "w:cs",
                "w:em",
                "w:lang",
                "w:eastAsianLayout",
                "w:specVanish",
                "w:oMath",
            ),
        )
        return spacing

    rPr = run._r.get_or_add_rPr()
    spacing = get_or_add_spacing(rPr)
    spacing.set(qn('w:val'), str(value))
    
    
# Iterate through rows in Excel and add to the Word document
def main(filename):
	start_time = time.time()
	workbook = load_workbook(filename, data_only=True)
	sheet = workbook.active

	doc = Document()
	doc.styles['Normal'].font.name = 'Times New Roman'
	doc.styles['Normal'].font.size = Pt(14)

	
	table = doc.add_table(rows=1, cols=1)
	table.style = 'Table Grid'
	cell = table.cell(0, 0)
	cell.text = str(sheet['B1'].value)
	cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
	doc.add_paragraph()
	
	for row in sheet.iter_rows(min_row=3, values_only=True):
	    if row[0] is None:
	        continue
	    paragraph = doc.add_paragraph()
	    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
	
	    for i, cell_value in enumerate(row):
	        if i == 1 or i == 2:
	            cell_value = clean_end_of_line(str(cell_value), add_dot=True)
	        else:
	            cell_value = clean_end_of_line(str(cell_value), add_dot=False)

	        cell_value = cell_value if cell_value is not None else ""
	        if i == 2:
	            capitalized_value = str(cell_value).capitalize()
	        else:
	            capitalized_value = str(cell_value)
	           
	        run = paragraph.add_run(capitalized_value)

	        		        
	        if i == 0:
	            run.font.bold = True
	            run.font.all_caps = True  
	        elif i == 1:
	            run_set_spacing(run, 60)
	        elif i == 2:
	            run.italic = True 

	          

	doc.save(filename.replace(file_extension, ".docx"))
	end_time = time.time()
	execution_time = round(end_time - start_time, 2)
	print(f"Скрипт жұмысы(парсинг, компиляция): {execution_time} секунд")
	
	
if __name__ == "__main__":
    if len(sys.argv) > 1:
        filename = sys.argv[1]
        print(f"Processing: {filename}")
        main(filename)
    else:
        print("Give some XLSX file")